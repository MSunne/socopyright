"""申请表.docx 渲染器。

基于 templates/application_form_template.docx（原始模板副本）按单元格坐标写入字段。
所有勾选框用 ☑/☐ 文本替换。手抄/签字/盖章区域保持空白。
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document

from .docx_utils import checkbox_line, set_cell_text

TEMPLATE = Path(__file__).resolve().parent.parent.parent / "templates" / "application_form_template.docx"


def _fmt_date_cn(d: str | date) -> str:
    if isinstance(d, date):
        return f"{d.year}年{d.month:02d}月{d.day:02d}日"
    # ISO "YYYY-MM-DD"
    y, m, dd = d.split("-")
    return f"{int(y)}年{int(m):02d}月{int(dd):02d}日"


def _hw_line(hw: dict) -> str:
    return f"CPU：{hw.get('cpu', '')}；内存：{hw.get('ram', '')}；硬盘：{hw.get('disk', '')}"


def render(spec: dict, *, output_path: str | Path, today: date | None = None) -> Path:
    """把 ProjectSpec 渲染成申请表.docx 并保存到 output_path。

    返回生成文件的 Path。
    """
    today = today or date.today()
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(TEMPLATE))

    owner = spec.get("owner", {})
    sw_name = spec["software_name"]
    version = spec.get("version", "V1.0")

    # ==== Table 0: 基本信息 + 著作权人 ====
    t0 = doc.tables[0]
    set_cell_text(t0.cell(0, 3), sw_name)           # 软件全称
    set_cell_text(t0.cell(0, 10), version)          # 版本号
    set_cell_text(t0.cell(1, 3), spec.get("software_abbr", ""))    # 软件简称
    set_cell_text(t0.cell(1, 10), spec.get("software_category", "应用软件"))  # 软件分类

    # 软件作品说明（勾选框）：spec 可配 is_original=True/False
    is_original = bool(spec.get("is_original", True))
    set_cell_text(
        t0.cell(2, 3),
        checkbox_line([("原创", is_original), ("修改（含翻译软件、合成软件）", not is_original)]),
    )
    # 开发完成日期
    set_cell_text(t0.cell(3, 3), _fmt_date_cn(spec["completion_date"]))
    # 发表状态：spec 可配 publish_status="已发表"/"未发表"
    is_published = (spec.get("publish_status", "未发表") == "已发表")
    set_cell_text(
        t0.cell(4, 3),
        checkbox_line([("已发表", is_published), ("未发表", not is_published)]),
    )
    # 开发方式：spec 可配 dev_mode（单独开发/合作开发/委托开发/下达任务开发）
    dev_mode = spec.get("dev_mode", "单独开发")
    set_cell_text(
        t0.cell(5, 3),
        checkbox_line([
            ("单独开发", dev_mode == "单独开发"),
            ("合作开发", dev_mode == "合作开发"),
            ("委托开发", dev_mode == "委托开发"),
            ("下达任务开发", dev_mode == "下达任务开发"),
        ]),
    )

    # 著作权人行
    set_cell_text(t0.cell(7, 1), owner.get("name", ""))  # 名称
    set_cell_text(t0.cell(7, 4), owner.get("type", "企业法人"))       # 类别
    set_cell_text(t0.cell(7, 5), owner.get("cert_type", "统一社会信用代码证书"))  # 证件类型
    set_cell_text(t0.cell(7, 6), owner.get("uscc", ""))               # 证件号
    set_cell_text(t0.cell(7, 7), owner.get("nationality", "中国"))    # 国籍
    prov_city = f"{owner.get('province', '')} {owner.get('city', '')}".strip()
    set_cell_text(t0.cell(7, 9), prov_city)                            # 省份/城市
    set_cell_text(t0.cell(7, 11), _fmt_date_cn(owner.get("established_date", today)))  # 成立日期

    # ==== Table 1: 权利说明 + 鉴别材料 + 软件功能 ====
    t1 = doc.tables[1]
    # 权利取得方式
    set_cell_text(
        t1.cell(0, 2),
        checkbox_line([("原始取得", True), ("继受取得", False)])
        + "\n" + checkbox_line([("该软件已登记", False), ("原登记做过变更或补充", False)]),
    )
    # 权利范围
    set_cell_text(t1.cell(1, 2), checkbox_line([("全部", True), ("部分", False)]))
    # 程序鉴别材料
    set_cell_text(
        t1.cell(2, 2),
        "☑ 一般交存：提交源程序前连续的30页和后连续的30页\n"
        "☐ 例外交存（使用黑色宽斜线覆盖 / 前10页和任选连续50页 / 目标程序前后各30页+源程序任选20页）",
    )
    # 文档鉴别材料
    set_cell_text(
        t1.cell(3, 2),
        "☑ 一般交存：提交任何一种文档的前连续的30页和后连续的30页\n"
        "☐ 例外交存（使用黑色宽斜线覆盖 / 前10页和任选连续50页）",
    )
    # 硬件环境
    set_cell_text(
        t1.cell(4, 2),
        f"开发：{_hw_line(spec['hardware_dev'])}\n运行：{_hw_line(spec['hardware_run'])}",
    )
    # 软件环境
    set_cell_text(
        t1.cell(5, 2),
        f"开发：{spec['dev_os']}；IDE：{spec['ide']}；数据库：{spec['database']}\n"
        f"运行：{spec['run_os']}；WEB容器：{spec['web_server']}；数据库：{spec['database']}",
    )
    # 编程语言
    set_cell_text(t1.cell(6, 2), " ; ".join(spec.get("language_list") or [spec.get("language", "")]))
    # 源程序量（Phase 3 源代码生成后回填，此处先显示 spec 里的值）
    set_cell_text(t1.cell(6, 4), f"{spec.get('source_lines', 0)}行")

    # 主要功能和技术特点（声明性视角：业务背景 + 行业 + 技术分类）
    # N10：只取 main_description 的第 1 段（业务背景），避免与"功能特点.docx"
    # 大段重复让审核员一眼看出复制粘贴感。
    md_paragraphs = [p.strip() for p in (spec.get("main_description", "") or "").split("\n\n") if p.strip()]
    md_first = md_paragraphs[0] if md_paragraphs else ""
    main_blob = (
        f"开发目的：{spec.get('purpose', '')}\n\n"
        f"面向领域／行业：{spec.get('industry', '')}\n\n"
        f"业务背景：{md_first}\n\n"
        f"技术特点分类：{spec.get('tech_category', '')}\n"
        f"技术特点：{spec.get('tech_features', '')}"
    )
    set_cell_text(t1.cell(7, 2), main_blob)

    # ==== Table 2: 申请办理 ====
    t2 = doc.tables[2]
    set_cell_text(
        t2.cell(0, 1),
        checkbox_line([("由著作权人申请", True), ("由代理人申请", False)]),
    )
    set_cell_text(t2.cell(1, 1), owner.get("name", ""))
    # t2[3,0] 代理人姓名：留空

    # t2[4,0/1] 手抄声明区：留空（用户必须手写）
    # 注意：模板里这块已经打印了声明内容 + 下方手抄区；我们保持该段落不动（仅要求不要填手抄区）
    # 因此这里不操作。

    # ==== Table 3: 证书份数 + 提交材料清单 ====
    t3 = doc.tables[3]
    set_cell_text(t3.cell(0, 1), "1 份正本     0 份副本")
    # 申请表页数固定 4 页
    set_cell_text(t3.cell(4, 1), "打印签字或盖章的登记申请表                                    4页")
    # 身份证明
    set_cell_text(
        t3.cell(5, 1),
        f"著作权人 - {owner.get('name', '')} 的{owner.get('cert_type', '统一社会信用代码证书')}复印件                                       1页",
    )
    # 软件鉴别材料页数：申请表 t1[2,2]/t1[3,2] 已声明"一般交存（前30+后30）"，
    # 此处填实际提交给版权中心的页数 = 60。pipeline 在打 zip 前会用
    # _clip_general_deposit() 把源代码/手册 PDF 截取为前 30 + 后 30 页（O0）。
    # 完整 PDF（仅做内部留档）保存在 _full/，不进 zip。
    set_cell_text(t3.cell(7, 1), "程序鉴别材料 - 一般交存\n文档鉴别材料 - 一般交存")
    set_cell_text(t3.cell(7, 2), "60页\n60页")

    # 签字日期（模板原有 "2026 年 04 月 13 日" 在正文段落里，不在表格里；保持原样）
    # 经办人签名 / 盖章 / 身份证号：均在正文段落，保持空白

    doc.save(str(output_path))
    return output_path
