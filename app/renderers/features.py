"""功能特点.docx 渲染器。

结构比申请表简单，只是一组 label / value 单元格。
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document

from .docx_utils import set_cell_text

TEMPLATE = Path(__file__).resolve().parent.parent.parent / "templates" / "features_template.docx"


def _fmt_date_cn(d: str | date) -> str:
    if isinstance(d, date):
        return f"{d.year}年{d.month:02d}月{d.day:02d}日"
    y, m, dd = d.split("-")
    return f"{int(y)}年{int(m):02d}月{int(dd):02d}日"


def _hw_line(hw: dict) -> str:
    return f"CPU：{hw.get('cpu', '')}；内存：{hw.get('ram', '')}；硬盘：{hw.get('disk', '')}"


def render(spec: dict, *, output_path: str | Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(TEMPLATE))
    owner = spec.get("owner", {})

    # ==== Table 0 ====
    t0 = doc.tables[0]
    set_cell_text(t0.cell(1, 1), spec["software_name"])
    set_cell_text(t0.cell(1, 3), spec.get("version", "V1.0"))
    set_cell_text(t0.cell(2, 1), spec.get("software_abbr", ""))
    set_cell_text(t0.cell(2, 3), spec.get("software_category", "应用软件"))
    # O7: 勾选项联动 spec（原创 / 开发方式 / 发表状态）
    set_cell_text(t0.cell(3, 1), "原创" if spec.get("is_original", True) else "修改")
    set_cell_text(t0.cell(3, 3), spec.get("dev_mode", "单独开发"))
    set_cell_text(t0.cell(4, 1), _fmt_date_cn(spec["completion_date"]))
    set_cell_text(t0.cell(4, 3), spec.get("publish_status", "未发表"))
    set_cell_text(t0.cell(5, 1), _hw_line(spec["hardware_dev"]))
    set_cell_text(t0.cell(6, 1), _hw_line(spec["hardware_run"]))
    set_cell_text(t0.cell(7, 1), spec.get("dev_os", ""))
    set_cell_text(t0.cell(8, 1), f"IDE：{spec.get('ide', '')}；数据库：{spec.get('database', '')}")
    set_cell_text(t0.cell(9, 1), spec.get("run_os", ""))
    set_cell_text(t0.cell(10, 1), f"WEB容器：{spec.get('web_server', '')}；数据库：{spec.get('database', '')}")
    set_cell_text(t0.cell(11, 1), "、".join(spec.get("language_list") or [spec.get("language", "")]))
    set_cell_text(t0.cell(12, 1), str(spec.get("source_lines", 0)))
    set_cell_text(t0.cell(13, 1), spec.get("purpose", ""))
    set_cell_text(t0.cell(14, 1), spec.get("industry", ""))

    # 主要功能：实施性视角——技术方法（main_description 第 2 段）+ 价值（第 3 段）+ 模块清单
    # N10：避免与申请表 t1[7,2] 重复，两份文档各取 main_description 不同段落
    md_paragraphs = [p.strip() for p in (spec.get("main_description", "") or "").split("\n\n") if p.strip()]
    md_method = md_paragraphs[1] if len(md_paragraphs) >= 2 else (md_paragraphs[0] if md_paragraphs else "")
    md_value = md_paragraphs[2] if len(md_paragraphs) >= 3 else ""
    func_lines = "\n".join(f"{i+1}. {f['name']}：{f['desc']}" for i, f in enumerate(spec.get("functions", [])))
    main_text = (
        f"核心方法：{md_method}\n\n"
        + (f"实际效果：{md_value}\n\n" if md_value else "")
        + "主要功能模块：\n" + func_lines
    )
    set_cell_text(t0.cell(15, 1), main_text)

    set_cell_text(t0.cell(16, 1), spec.get("tech_category", ""))
    set_cell_text(t0.cell(17, 1), spec.get("tech_features", ""))

    # ==== Table 1: 著作权人 ====
    # O9: type/cert_type 联动 USCC 第 1 位（spec.generate_specs 已经填好 owner.type 和 cert_type）
    t1 = doc.tables[1]
    set_cell_text(t1.cell(0, 1), owner.get("name", ""))
    set_cell_text(t1.cell(1, 1), owner.get("type", "企业法人"))
    set_cell_text(t1.cell(2, 1), owner.get("cert_type", "统一社会信用代码证书"))
    set_cell_text(t1.cell(3, 1), owner.get("uscc", ""))
    set_cell_text(t1.cell(4, 1), f"{owner.get('province', '')} {owner.get('city', '')}".strip())

    doc.save(str(output_path))
    return output_path
